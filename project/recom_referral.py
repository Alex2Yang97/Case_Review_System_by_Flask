import io
import os
import re
import docx
import pandas as pd
from datetime import datetime
from dateutil.relativedelta import relativedelta
from flask import render_template, url_for, flash, redirect, Blueprint, make_response, Response
from flask_login import login_required, current_user
from wtforms.validators import InputRequired, Optional

from werkzeug.exceptions import abort
from werkzeug.utils import secure_filename

from .models import get_user_info
from .forms import (UploadWordReport, UploadRecom, UploadSAR, UploadSanc,
                    FilterRecom, RecomendationForm, AddRecomendationForm, 
                    FilterSAR, SARForm, AddSARForm,
                    FilterSanc, SanctionForm, AddSanctionForm)
from .utils import connect_db, execute_sql, select_sql, get_db_engine, show_step_card, get_case, cal_dates, strptime_date
from .config import EDD_ANS_REPORT_DIR, SQL_TYPE




recom_referral = Blueprint('recom_referral', __name__)


def read_recom(filename, case_id):
    doc = docx.Document(filename)
    for i, para in enumerate(doc.paragraphs):
        if "Consolidated Recommendation" in para.text:
            start = i+1
        if "Relationship Overview" in para.text:
            end = i
            break
    recom_parts = doc.paragraphs[start: end]
    
    from_section_lst = []
    escal_recom_details = []
    detail = []
    for recom in recom_parts:
        if recom.text != "":
            section = 0
            for run in recom.runs:
                if run.underline:
                    from_section_lst.append(recom.text.replace("\t", " "))
                    section = 1
                    if detail:
                        escal_recom_details.append(" ".join(detail))
                        detail = []
                    break
            if section == 0:
                detail.append(recom.text)
    if detail:
        escal_recom_details.append(" ".join(detail))
        
    recom_data = pd.DataFrame({
        "Case_ID": case_id,
        "Status": "Open",
        "From_Section": from_section_lst,
        "Escal_Recomm_Details": escal_recom_details
    })
    return recom_data


def read_sanc(filename, case_id):
    doc = docx.Document(filename)
    for i, para in enumerate(doc.paragraphs):
        if ("Sanctions Referral" in para.text) or ("Sanctions Referrals" in para.text):
            start = i+1
        if "Report Sign-Offs" in para.text:
            end = i
            break
    sanc_parts = doc.paragraphs[start: end]
    
    # 去掉加粗的部分（subject_name均是有下划线且未加粗）
    filtered_sanc_parts = []
    for para in sanc_parts:
        bold = 0
        for run in para.runs:
            if run.bold:
                bold = 1
        if bold == 0:
            filtered_sanc_parts.append(para)
            
    sub_name_lst = []
    referral_reason_lst = []
    detail = []
    for sanc in filtered_sanc_parts:
        if sanc.text != "":
            underline = 0
            for run in sanc.runs:
                if run.underline:
                    sub_name_lst.append(sanc.text)
                    if detail:
                        referral_reason_lst.append(" ".join(detail))
                        detail = []
                    underline = 1
                    break
            if underline == 0:
                detail.append(sanc.text)
    if detail:
        referral_reason_lst.append(" ".join(detail))

    sanc_data = pd.DataFrame({
        "Case_ID": case_id,
        "Status": "Open",
        "Subject_Name": sub_name_lst,
        "Referral_Reason": referral_reason_lst
    })
    return sanc_data


def read_sar(filename, case_id):
    doc = docx.Document(filename)
    for i, para in enumerate(doc.paragraphs):
        if ("SAR Referral" in para.text) or ("SAR Referrals" in para.text):
            start = i+1
        if ("Report Sign-Offs" in para.text) or ("Sanctions Referral" in para.text):
            end = i
            break
    sar_parts = doc.paragraphs[start: end]
    
    # Get each Entity and Details
    acc_lst = []
    acc_index_lst = []
    acc_detail_lst = []
    for i, para in enumerate(sar_parts):
        if para.text and "account" in para.text.lower():
            acc_lst.append(para.text)
            acc_index_lst.append(i)
            
    details = ""
    for i, para in enumerate(sar_parts):
        if i in acc_index_lst:
            if details:
                acc_detail_lst.append(details)
                details = ""
        else:
            bold = 0
            for run in para.runs:
                if run.bold:
                    bold = 1
                    break
            if bold == 0:
                details += para.text
    acc_detail_lst.append(details)

    # Process each Entity and Details
    sub_name_lst = []
    sub_acc_lst = []
    amount_lst = []
    start_date_lst = []
    end_date_lst = []

    for acc_name, acc_detail in zip(acc_lst, acc_detail_lst):

        # Get amount
        dollar_regex = r"\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?"
        amounts = re.findall(dollar_regex, acc_detail)
        if amounts:
            amount_lst.append(float(amounts[0].replace("$", "", 1).replace(",", "")))
        else:
            amount_lst.append(0)
        
        # Get dates
        date_regex = r"\d{1,2}/\d{1,2}/\d{4}"
        dates = re.findall(date_regex, acc_detail)
        dates = sorted([datetime.strptime(date_str, "%m/%d/%Y") for date_str in dates])
        dates = [dt.strftime("%m/%d/%Y") for dt in dates]
        start_date_lst.append(dates[0])
        end_date_lst.append(dates[-1])
        
        # Get accounts
        acc_lst = []
        account_regex = r"#(.*?)( |\))"
        matches = re.findall(account_regex, acc_name)
        for match in matches:
            acc_lst.append(match[0])
        sub_acc_lst.append(",".join(acc_lst))
        
        # Get subject names
        name_regex = r"^(.*?)\s*\("
        match = re.match(name_regex, acc_name)
        if match:
            sub_name_lst.append(match.group(1))
        
    sar_data = pd.DataFrame({
        "Case_ID": case_id,
        "Status": "Open",
        "Subject_Name": sub_name_lst,
        "Subject_Account_NO": sub_acc_lst,
        "Activity_Start_Date": start_date_lst,
        "Activity_End_Date": end_date_lst,
        "Amount": amount_lst,
        "Referral_Reason": acc_detail_lst
    })
    return sar_data


def get_work_progress(case_id):
    cursor, cnxn = connect_db()
    work_progress = pd.read_sql(f"select Recommendation_ID, Status from Recommendation_local where Case_ID = {case_id} and Status <> 'Removed' ", cnxn)
    total_recoms = work_progress["Recommendation_ID"].tolist()
    open_recoms = work_progress[work_progress["Status"]=="Open"]["Recommendation_ID"].tolist()
    closed_recoms = work_progress[work_progress["Status"]=="Closed"]["Recommendation_ID"].tolist()
    removed_recoms = work_progress[work_progress["Status"]=="Removed"]["Recommendation_ID"].tolist()
    
    work_progress = pd.read_sql(f"select SAR_Referral_ID, Status from SARref_local where Case_ID = {case_id} and Status <> 'Removed' ", cnxn)
    total_sars = work_progress["SAR_Referral_ID"].tolist()
    open_sars = work_progress[work_progress["Status"]=="Open"]["SAR_Referral_ID"].tolist()
    closed_sars = work_progress[work_progress["Status"]=="Closed"]["SAR_Referral_ID"].tolist()
    removed_sars = work_progress[work_progress["Status"]=="Removed"]["SAR_Referral_ID"].tolist()
    
    work_progress = pd.read_sql(f"select Sanction_Referral_ID, Status from SanctionRef_local where Case_ID = {case_id} and Status <> 'Removed' ", cnxn)
    total_sancs = work_progress["Sanction_Referral_ID"].tolist()
    open_sancs = work_progress[work_progress["Status"]=="Open"]["Sanction_Referral_ID"].tolist()
    closed_sancs = work_progress[work_progress["Status"]=="Closed"]["Sanction_Referral_ID"].tolist()
    removed_sancs = work_progress[work_progress["Status"]=="Removed"]["Sanction_Referral_ID"].tolist()
    cursor.close()
    
    return (total_recoms, open_recoms, closed_recoms, removed_recoms, 
            total_sars, open_sars, closed_sars, removed_sars,
            total_sancs, open_sancs, closed_sancs, removed_sancs)


def close_case(case_id):
    sql_select = f'''
        select FID_KYC_Refresh_Date, Risk_Rating, Escal_Recom_Status, SAR_Referral_Status, Sanction_Referral_Status, Customer_ID, 
        Customer_Name, Risk_Rating, Type, Category, Case_Type, Comment from casetracking_local where case_id = {case_id}'''    
    case_info = select_sql(sql_select)
    
    if case_info["Escal_Recom_Status"] == "Closed" and case_info["SAR_Referral_Status"] == "Closed" and case_info["Sanction_Referral_Status"] == "Closed":
        sql_update = f"update CaseTracking_local SET Case_Status = 'Closed', Sub_Status = 'Completed' where Case_ID = {case_id}"
        execute_sql(sql_update)
        flash("All steps has been completed!", "success")
    return
    

@recom_referral.route('/upload/<int:case_id>', methods=["GET", "POST"])
@login_required
def upload_page(case_id):
    upload_word = UploadWordReport()
    upload_recom = UploadRecom()
    upload_sar = UploadSAR()
    upload_sanc = UploadSanc()
    
    if upload_recom.clear_recom_btn.data and upload_recom.validate():
        sql_clear = f'''
            update CaseTracking_local SET 
                Case_Status = 'Recommendation & Referral', Sub_Status = 'Approved', 
                Escal_Recom_Status = NULL, Has_Recommendation = NULL where Case_ID = {case_id}'''
        execute_sql(sql_clear)
        sql_clear = f"DELETE FROM Recommendation_local WHERE case_id = {case_id} "
        execute_sql(sql_clear)
        flash(f"Clear all Recommendations for {case_id}", "warning")
        return redirect(url_for("recom_referral.upload_page", case_id=case_id))
    
    if upload_sar.clear_sar_btn.data and upload_recom.validate():
        sql_clear = f'''
            update CaseTracking_local SET 
                Case_Status = 'Recommendation & Referral', Sub_Status = 'Approved', 
                SAR_Referral_Status = NULL, Has_SAR_Referral = NULL where Case_ID = {case_id}'''
        execute_sql(sql_clear)
        sql_clear = f"DELETE FROM SARref_local WHERE case_id = {case_id} "
        execute_sql(sql_clear)
        flash(f"Clear all SAR Referrals for {case_id}", "warning")
        return redirect(url_for("recom_referral.upload_page", case_id=case_id))
    
    if upload_sanc.clear_sanc_btn.data and upload_recom.validate():
        sql_clear = f'''
            update CaseTracking_local SET 
                Case_Status = 'Recommendation & Referral', Sub_Status = 'Approved', 
                Sanction_Referral_Status = NULL, Has_Sanction_Referral = NULL where Case_ID = {case_id}'''
        execute_sql(sql_clear)
        sql_clear = f"DELETE FROM SanctionRef_local WHERE case_id = {case_id} "
        execute_sql(sql_clear)
        flash(f"Clear all Sanction Referrals for {case_id}", "warning")
        return redirect(url_for("recom_referral.upload_page", case_id=case_id))
    
    if upload_recom.reopen_recom_btn.data and upload_recom.validate():
        sql_reopen = f'''
            update CaseTracking_local SET 
                Case_Status = 'Recommendation & Referral', Sub_Status = 'Approved', 
                Escal_Recom_Status = 'Open', Has_Recommendation = NULL where Case_ID = {case_id}'''
        execute_sql(sql_reopen)
        sql_reopen = f"update Recommendation_local SET Status = 'Open' where Case_ID = {case_id} and Status <> 'Removed' "
        execute_sql(sql_reopen)
        flash(f"Reopen all Recommendations for {case_id}", "success")
        return redirect(url_for("recom_referral.upload_page", case_id=case_id))
        
    if upload_sar.reopen_sar_btn.data and upload_recom.validate():
        sql_reopen = f'''
            update CaseTracking_local SET 
                Case_Status = 'Recommendation & Referral', Sub_Status = 'Approved', 
                SAR_Referral_Status = 'Open', Has_SAR_Referral = NULL where Case_ID = {case_id}'''
        execute_sql(sql_reopen)
        sql_reopen = f"update SARref_local SET Status = 'Open' where Case_ID = {case_id} and Status <> 'Removed' "
        execute_sql(sql_reopen)
        flash(f"Reopen all SAR Referrals for {case_id}", "success")
        return redirect(url_for("recom_referral.upload_page", case_id=case_id))
    
    if upload_sanc.reopen_sanc_btn.data and upload_recom.validate():
        sql_reopen = f'''
            update CaseTracking_local SET 
                Case_Status = 'Recommendation & Referral', Sub_Status = 'Approved', 
                Sanction_Referral_Status = 'Open', Has_Sanction_Referral = NULL where Case_ID = {case_id}'''
        execute_sql(sql_reopen)
        sql_reopen = f"update SanctionRef_local SET Status = 'Open' where Case_ID = {case_id} and Status <> 'Removed' "
        execute_sql(sql_reopen)
        flash(f"Reopen all Sanction Referrals for {case_id}", "success")
        return redirect(url_for("recom_referral.upload_page", case_id=case_id))
    
    # Upload word file
    if upload_word.upload_word_btn.data and upload_word.validate():
        if not upload_word.word_report.data:
            flash("Please upload EDD Report!", "danger")
        else:
            uploaded_file = upload_word.word_report.data
            
            # Save file in local
            filename = secure_filename(uploaded_file.filename)
            saved_file_dir = os.path.join(EDD_ANS_REPORT_DIR, f"user_id-{current_user.get_id()}")
            if not os.path.exists(saved_file_dir):
                os.makedirs(saved_file_dir)
                
            # uploaded_file.save(os.path.join(saved_file_dir, make_unique(filename)))
            timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            saved_file_name = os.path.join(saved_file_dir, timestamp+"_"+filename)
            uploaded_file.save(saved_file_name)
            
            message = []
            delete_doc = True
            
            try:
                recom_data = read_recom(uploaded_file, case_id)
                # 如果没有在文档中检测到Recommendation/Referral，最好让用户自己决定是否Close
                if not recom_data.empty:
                    if SQL_TYPE == "sql server":
                        engine = get_db_engine()
                        recom_data.to_sql(name='Recommendation_local', schema='dbo', con=engine, if_exists='append', index=False)
                    else:
                        cursor, cnxn = connect_db()
                        recom_data.to_sql(name='Recommendation_local', if_exists='append', con=cnxn, index=False)
                        cursor.close()
                        cnxn.close()
                        
                    update_sql = f"update CaseTracking_local SET Has_Recommendation = 'Yes', Escal_Recom_Status = 'Open' where Case_ID = {case_id}"
                    execute_sql(update_sql)
                
                    message.append("Recommendation")
                    delete_doc = False
            except:
                pass
            
            try:
                sar_data = read_sar(uploaded_file, case_id)
                if not sar_data.empty:                    
                    if SQL_TYPE == "sql server":
                        engine = get_db_engine()
                        sar_data.to_sql(name='SARref_local', schema='dbo', con=engine, if_exists='append', index=False)
                    else:
                        cursor, cnxn = connect_db()
                        sar_data.to_sql(name='SARref_local', if_exists='append', con=cnxn, index=False)
                        cursor.close()
                        cnxn.close()
                        
                    update_sql = f"update CaseTracking_local SET Has_SAR_Referral = 'Yes', SAR_Referral_Status = 'Open' where Case_ID = {case_id}"
                    execute_sql(update_sql)
                
                    message.append("SAR Referral")
                    delete_doc = False
            except:
                pass
            
            try:
                sanc_data = read_sanc(uploaded_file, case_id)                    
                if not sanc_data.empty:
                    update_sql = f"update CaseTracking_local SET Has_Sanction_Referral = 'Yes', Sanction_Referral_Status = 'Open' where Case_ID = {case_id}"
                    execute_sql(update_sql)
                    if SQL_TYPE == "sql server":
                        engine = get_db_engine()
                        sanc_data.to_sql(name='SanctionRef_local', schema='dbo', con=engine, if_exists='append', index=False)
                    else:
                        cursor, cnxn = connect_db()
                        sanc_data.to_sql(name='SanctionRef_local', if_exists='append', con=cnxn, index=False)
                        cursor.close()
                        cnxn.close()
                        
                    message.append("Sanction Referral")
                    delete_doc = False
            except:
                pass
        
        if delete_doc:
            if os.path.exists(saved_file_name):
                os.remove(saved_file_name)
            flash("Extract no record from document! Please check your document or add record manually!", "danger")
        else:
            message = ",".join(message)
            flash(f"Extract message from document!", "success")
        
        return redirect(url_for("recom_referral.upload_page", case_id=case_id))
    
    if upload_recom.add_recom_btn.data and upload_recom.validate():
        return redirect(url_for('recom_referral.add_recom', case_id=case_id))
    
    if upload_recom.no_recom_btn.data and upload_recom.validate():
        sql_update = f"update CaseTracking_local SET Has_Recommendation = 'No', Escal_Recom_Status = 'Closed' where Case_ID = {case_id}"
        execute_sql(sql_update)
        close_case(case_id)
        return redirect(url_for('recom_referral.upload_page', case_id=case_id))
        
    if upload_recom.check_recom_btn.data and upload_recom.validate():
        return redirect(url_for('recom_referral.preview_recom', case_id=case_id)) 
        
    if upload_recom.finish_recom_btn.data and upload_recom.validate():
        sql_close = f"update CaseTracking_local SET Escal_Recom_Status = 'Closed' where Case_ID = {case_id}"
        execute_sql(sql_close)
        sql_close = f"update Recommendation_local SET Status = 'Closed' where Case_ID = {case_id} and Status <> 'Removed' "
        execute_sql(sql_close)
        close_case(case_id)
        return redirect(url_for('recom_referral.upload_page', case_id=case_id))
    
    if upload_sar.add_sar_btn.data and upload_sar.validate():
        return redirect(url_for('recom_referral.add_sar', case_id=case_id))
    
    if upload_sar.no_sar_btn.data and upload_sar.validate():
        sql_update = f"update CaseTracking_local SET Has_SAR_Referral = 'No', SAR_Referral_Status = 'Closed' where Case_ID = {case_id}"
        execute_sql(sql_update)
        close_case(case_id)
        return redirect(url_for('recom_referral.upload_page', case_id=case_id))
    
    if upload_sar.check_sar_btn.data and upload_sar.validate():
        return redirect(url_for('recom_referral.preview_sar', case_id=case_id)) 
        
    if upload_sar.finish_sar_btn.data and upload_sar.validate():
        sql_close = f"update CaseTracking_local SET SAR_Referral_Status = 'Closed' where Case_ID = {case_id}"
        execute_sql(sql_close)
        sql_close = f"update SARref_local SET Status = 'Closed' where Case_ID = {case_id} and Status <> 'Removed' "
        execute_sql(sql_close)
        close_case(case_id)
        return redirect(url_for('recom_referral.upload_page', case_id=case_id))

    if upload_sanc.add_sanc_btn.data and upload_sanc.validate():
        return redirect(url_for('recom_referral.add_sanc', case_id=case_id))

    if upload_sanc.no_sanc_btn.data and upload_sanc.validate():
        sql_update = f"update CaseTracking_local SET Has_Sanction_Referral = 'No', Sanction_Referral_Status = 'Closed' where Case_ID = {case_id}"
        execute_sql(sql_update)
        close_case(case_id)
        return redirect(url_for('recom_referral.upload_page', case_id=case_id))
    
    if upload_sanc.check_sanc_btn.data and upload_sanc.validate():
        return redirect(url_for('recom_referral.preview_sanc', case_id=case_id)) 
        
    if upload_sanc.finish_sanc_btn.data and upload_sanc.validate():
        sql_close = f"update CaseTracking_local SET Sanction_Referral_Status = 'Closed' where Case_ID = {case_id}"
        execute_sql(sql_close)
        sql_close = f"update SanctionRef_local SET Status = 'Closed' where Case_ID = {case_id} and Status <> 'Removed' "
        execute_sql(sql_close)
        close_case(case_id)
        return redirect(url_for('recom_referral.upload_page', case_id=case_id)) 
    
    if get_user_info(current_user.get_id()).report != 1:
        for filed in upload_word:
            filed.render_kw = {"disabled": "disabled"}
        for filed in upload_recom:
            filed.render_kw = {"disabled": "disabled"}
        for filed in upload_sar:
            filed.render_kw = {"disabled": "disabled"}
        for filed in upload_sanc:
            filed.render_kw = {"disabled": "disabled"}
    
    sql_select = f'''
        select Case_Status, Sub_Status, Escal_Recom_Status, SAR_Referral_Status, Sanction_Referral_Status from casetracking_local where case_id = {case_id}'''
    case_info_default = select_sql(sql_select)
    
    if case_info_default["Case_Status"] != "Recommendation & Referral":
        for filed in upload_word:
            filed.render_kw = {"disabled": "disabled"}
        upload_recom.add_recom_btn.render_kw = {"disabled": "disabled"}
        upload_recom.no_recom_btn.render_kw = {"disabled": "disabled"}
        upload_sar.add_sar_btn.render_kw = {"disabled": "disabled"}
        upload_sar.no_sar_btn.render_kw = {"disabled": "disabled"}
        upload_sanc.add_sanc_btn.render_kw = {"disabled": "disabled"}
        upload_sanc.no_sanc_btn.render_kw = {"disabled": "disabled"}
            
    if not case_info_default["Escal_Recom_Status"]:
        upload_recom.check_recom_btn.render_kw = {"disabled": "disabled"}
        upload_recom.finish_recom_btn.render_kw = {"disabled": "disabled"}
        upload_recom.reopen_recom_btn.render_kw = {"disabled": "disabled"}
    else:
        upload_recom.no_recom_btn.render_kw = {"disabled": "disabled"}
        if case_info_default["Escal_Recom_Status"] == "Open":
            upload_recom.reopen_recom_btn.render_kw = {"disabled": "disabled"}
        if case_info_default["Escal_Recom_Status"] == "Closed":
            upload_recom.add_recom_btn.render_kw = {"disabled": "disabled"}
            upload_recom.finish_recom_btn.render_kw = {"disabled": "disabled"}
        
    if not case_info_default["SAR_Referral_Status"]:
        upload_sar.check_sar_btn.render_kw = {"disabled": "disabled"}
        upload_sar.finish_sar_btn.render_kw = {"disabled": "disabled"}
        upload_sar.reopen_sar_btn.render_kw = {"disabled": "disabled"}
    else:
        upload_sar.no_sar_btn.render_kw = {"disabled": "disabled"}
        if case_info_default["SAR_Referral_Status"] == "Open":
            upload_sar.reopen_sar_btn.render_kw = {"disabled": "disabled"}
        if case_info_default["SAR_Referral_Status"] == "Closed":
            upload_sar.add_sar_btn.render_kw = {"disabled": "disabled"}
            upload_sar.finish_sar_btn.render_kw = {"disabled": "disabled"}
        
    if not case_info_default["Sanction_Referral_Status"]:
        upload_sanc.check_sanc_btn.render_kw = {"disabled": "disabled"}
        upload_sanc.finish_sanc_btn.render_kw = {"disabled": "disabled"}
        upload_sanc.reopen_sanc_btn.render_kw = {"disabled": "disabled"}
    else:
        upload_sanc.no_sanc_btn.render_kw = {"disabled": "disabled"}
        if case_info_default["Sanction_Referral_Status"] == "Open":
            upload_sanc.reopen_sanc_btn.render_kw = {"disabled": "disabled"}
        if case_info_default["Sanction_Referral_Status"] == "Closed":
            upload_sanc.add_sanc_btn.render_kw = {"disabled": "disabled"}
            upload_sanc.finish_sanc_btn.render_kw = {"disabled": "disabled"}
    
    # Work Flow Chart
    steps = show_step_card(case_id, case_info_default["Case_Status"])
    
    (total_recoms, open_recoms, closed_recoms, _, 
     total_sars, open_sars, closed_sars, _,
     total_sancs, open_sancs, closed_sancs, _) = get_work_progress(case_id)
    
    return render_template(
        'upload_recom_referral.html', case_id=case_id, upload_word=upload_word,
        upload_recom=upload_recom, upload_sar=upload_sar, upload_sanc=upload_sanc,
        recom_status=case_info_default["Escal_Recom_Status"], 
        sar_status=case_info_default["SAR_Referral_Status"], 
        sanc_status=case_info_default["Sanction_Referral_Status"],
        total_recoms=total_recoms, open_recoms=open_recoms, closed_recoms=closed_recoms,
        total_sars=total_sars, open_sars=open_sars, closed_sars=closed_sars,
        total_sancs=total_sancs, open_sancs=open_sancs, closed_sancs=closed_sancs,
        steps=steps, case_status=case_info_default["Case_Status"], 
        sub_status=case_info_default["Sub_Status"],
        case=get_case(case_id)
        )

#============================= Working on Recommendation =====================================================
@recom_referral.route('/recom/<int:case_id>', methods=["GET", "POST"])
@login_required
def preview_recom(case_id):
    filter_recom_form = FilterRecom()
    
    cursor, cnxn = connect_db()
    work_progress = pd.read_sql(f"select Recommendation_ID, Status from Recommendation_local where Case_ID = {case_id} and Status <> 'Removed' ", cnxn)
    total_tasks = work_progress["Recommendation_ID"].tolist()
    open_tasks = work_progress[work_progress["Status"]=="Open"]["Recommendation_ID"].tolist()
    cursor.close()
    cnxn.close()
    
    if filter_recom_form.add_recom.data and filter_recom_form.validate():
        return redirect(url_for("recom_referral.add_recom", case_id=case_id))
    
    if filter_recom_form.filter_recom.data and filter_recom_form.validate():
        cursor, cnxn = connect_db()
        sql_get_recoms = f'''
            select Recommendation_ID, Status, Closure_Date, From_Section, Initiated_Date, Ack_Action_Date, Recomm_or_Escal, 
            Responsible_Personnel, Action_Details, Followup_Date_1st, Followup_Date_2nd, Last_Followup_Date, Escalation_Date, 
            Escalation_Type, Escalated_To, Recommendation_Closure_Details, Escal_Recomm_Details from Recommendation_local 
            where Case_ID = {case_id} and Status <> 'Removed' '''
        if filter_recom_form.recom_status.data != "ALL":
             sql_get_recoms += f" and Status = '{filter_recom_form.recom_status.data}' "
        recoms = pd.read_sql(sql_get_recoms, cnxn)
        cursor.close()
        cnxn.close()
        
        return render_template(
            'preview_case_recom.html', case_id=case_id, filter_recom_form=filter_recom_form, recoms=recoms, 
            total_tasks=total_tasks, open_tasks=open_tasks
            )
    
    if filter_recom_form.export_recom.data and filter_recom_form.validate():
        cursor, cnxn = connect_db()
        sql_get_recoms = (
            "select Case_ID, Recommendation_ID, Status, Closure_Date, From_Section, Initiated_Date, Ack_Action_Date, Recomm_or_Escal, "
            "Responsible_Personnel, Action_Details, Followup_Date_1st, Followup_Date_2nd, Last_Followup_Date, Escalation_Date, "
            "Escalation_Type, Escalated_To, Recommendation_Closure_Details, Escal_Recomm_Details from Recommendation_local "
            f"where Case_ID = {case_id} and Status <> 'Removed' ")
        if filter_recom_form.recom_status.data != "ALL":
             sql_get_recoms += f" and Status = '{filter_recom_form.recom_status.data}' "        
        
        recom_df = pd.read_sql(sql_get_recoms, cnxn)
        cursor.close()
        cnxn.close()
        
        # resp = make_response(recom_df.to_csv(index=False))
        # resp.headers["Content-Disposition"] = f"attachment; filename=Recom-CaseID_{case_id}-Status_{filter_recom_form.recom_status.data}.csv"
        # resp.headers["Content-Type"] = "text/csv"
        # return resp
    
        buffer = io.BytesIO()
        recom_df.to_excel(buffer, index=False)
        headers = {
            'Content-Disposition': f'attachment; filename=Recom-CaseID_{case_id}-Status_{filter_recom_form.recom_status.data}.xlsx',
            'Content-type': 'application/vnd.ms-excel'
        }
        return Response(buffer.getvalue(), mimetype='application/vnd.ms-excel', headers=headers)
    
    cursor, cnxn = connect_db()
    sql_get_recoms = f'''
        select Recommendation_ID, Status, Closure_Date, From_Section, Initiated_Date, Ack_Action_Date, Recomm_or_Escal, 
        Responsible_Personnel, Action_Details, Followup_Date_1st, Followup_Date_2nd, Last_Followup_Date, Escalation_Date, 
        Escalation_Type, Escalated_To, Recommendation_Closure_Details, Escal_Recomm_Details from Recommendation_local 
        where Case_ID = {case_id} and Status <> 'Removed' '''
    recoms = pd.read_sql(sql_get_recoms, cnxn)
    cursor.close()
    cnxn.close()
        
    return render_template(
        'preview_case_recom.html', case_id=case_id, filter_recom_form=filter_recom_form, recoms=recoms, 
        total_tasks=total_tasks, open_tasks=open_tasks
        )


@recom_referral.route('/recom/<int:case_id>/<int:recom_id>', methods=["GET", "POST"])
@login_required
def each_recom(case_id, recom_id):
    recom_form = RecomendationForm()
    
    if recom_form.remove_btn.data and recom_form.validate():
        sql_remove = f"update Recommendation_local set Status = 'Removed' where Recommendation_ID = {recom_id}"
        execute_sql(sql_remove)
        flash("Remove Recommendation!", "warning")
        return redirect(url_for('recom_referral.each_recom', case_id=case_id, recom_id=recom_id))
    
    if recom_form.close_btn.data and recom_form.validate():
        sql_close = f"update Recommendation_local set Status = 'Closed', Closure_Date = ? where Recommendation_ID = {recom_id}"
        execute_sql(sql_close, (datetime.now().date(),))
        flash("Close Recommendation successfully!", "success")
        return redirect(url_for('recom_referral.each_recom', case_id=case_id, recom_id=recom_id))
    
    if recom_form.export_btn.data and recom_form.validate():
        # Export excel to user computer
        cursor, cnxn = connect_db()
        sql_get_recom = (
            "select Case_ID, Recommendation_ID, Status, Initiated_Date, Closure_Date, From_Section, Initiated_Date, Ack_Action_Date, Recomm_or_Escal, "
            "Responsible_Personnel, Action_Details, Followup_Date_1st, Followup_Date_2nd, Last_Followup_Date, Escalation_Date, "
            "Escalation_Type, Escalated_To, Recommendation_Closure_Details, Escal_Recomm_Details from Recommendation_local "
            f"where Recommendation_ID = {recom_id}")
        recom_df = pd.read_sql(sql_get_recom, cnxn)
        cursor.close()
        cnxn.close()
        
        # resp = make_response(recom_df.to_csv(index=False))
        # resp.headers["Content-Disposition"] = f"attachment; filename=Case-{case_id}-recom-{recom_id}.csv"
        # resp.headers["Content-Type"] = "text/csv"
        # return resp
    
        buffer = io.BytesIO()
        recom_df.to_excel(buffer, index=False)
        headers = {
            'Content-Disposition': f'attachment; filename=Case-{case_id}-recom-{recom_id}.xlsx',
            'Content-type': 'application/vnd.ms-excel'
        }
        return Response(buffer.getvalue(), mimetype='application/vnd.ms-excel', headers=headers)
    
    if recom_form.update_btn.data and recom_form.validate():
        sql_update = (
            "update Recommendation_local set From_Section = ?, Ack_Action_Date = ?, Recomm_or_Escal = ?, Responsible_Personnel = ?, "
            "Action_Details = ?, Followup_Date_1st = ?, Followup_Date_2nd = ?, Last_Followup_Date = ?, Escalation_Date = ?, "
            "Escalation_Type = ?, Escalated_To = ?, Recommendation_Closure_Details = ?, Escal_Recomm_Details = ? "
            f"where Recommendation_ID = {recom_id}")
        execute_sql(sql_update, (
            recom_form.from_section.data, recom_form.ack_action_date.data, recom_form.recom_or_esc.data, recom_form.responsible_personnel.data, 
            recom_form.action_detail.data, recom_form.followup_1_date.data, recom_form.followup_2_date.data, recom_form.followup_last_date.data, recom_form.escalation_date.data,
            recom_form.escalation_type.data, recom_form.escalation_to.data, recom_form.closure_detail.data, recom_form.escalation_recom_detail.data
        ))
        flash("Update Recommendation successfully!", "success")
        return redirect(url_for('recom_referral.each_recom', case_id=case_id, recom_id=recom_id))
    
    sql_get_recom = f'''
        select Recommendation_ID, Status, Initiated_Date, Closure_Date, From_Section, Ack_Action_Date, Recomm_or_Escal, 
        Responsible_Personnel, Action_Details, Followup_Date_1st, Followup_Date_2nd, Last_Followup_Date, Escalation_Date, 
        Escalation_Type, Escalated_To, Recommendation_Closure_Details, Escal_Recomm_Details from Recommendation_local 
        where Recommendation_ID = {recom_id}'''
    recom_info = select_sql(sql_get_recom)
    
    recom_form.recom_status.default = recom_info["Status"]
    recom_form.initiated_date.default = strptime_date(recom_info["Initiated_Date"])
    recom_form.closure_date.default = strptime_date(recom_info["Closure_Date"])
    recom_form.recom_or_esc.default = recom_info["Recomm_or_Escal"]
    recom_form.responsible_personnel.default = recom_info["Responsible_Personnel"]
    recom_form.action_detail.default = recom_info["Action_Details"]
    recom_form.followup_1_date.default = strptime_date(recom_info["Followup_Date_1st"])
    recom_form.followup_2_date.default = strptime_date(recom_info["Followup_Date_2nd"])
    recom_form.followup_last_date.default = strptime_date(recom_info["Last_Followup_Date"])
    recom_form.escalation_date.default = strptime_date(recom_info["Escalation_Date"])
    recom_form.escalation_type.default = recom_info["Escalation_Type"]
    recom_form.escalation_to.default = recom_info["Escalated_To"]
    recom_form.closure_detail.default = recom_info["Recommendation_Closure_Details"]
    recom_form.from_section.default = recom_info["From_Section"]
    recom_form.ack_action_date.default = strptime_date(recom_info["Ack_Action_Date"])
    recom_form.escalation_recom_detail.default = recom_info["Escal_Recomm_Details"]
    recom_form.process()

    if recom_info["Status"] == "Removed" or recom_info["Status"] == "Closed":
        # If validator is InputRequired(), then "disabled" will make form.validate() always False
        for myform_field in recom_form:
            validators = myform_field.validators
            for i, validator in enumerate(validators):
                if isinstance(validator, InputRequired):
                    validators[i] = Optional()
        
        recom_form.ack_action_date.render_kw = {"disabled": "disabled"}
        recom_form.recom_or_esc.render_kw = {"disabled": "disabled"}
        recom_form.responsible_personnel.render_kw = {"disabled": "disabled"}
        recom_form.followup_1_date.render_kw = {"disabled": "disabled"}
        recom_form.followup_2_date.render_kw = {"disabled": "disabled"}
        recom_form.followup_last_date.render_kw = {"disabled": "disabled"}
        recom_form.escalation_date.render_kw = {"disabled": "disabled"}
        recom_form.escalation_type.render_kw = {"disabled": "disabled"}
        recom_form.escalation_to.render_kw = {"disabled": "disabled"}
        recom_form.from_section.render_kw = {"disabled": "disabled"}
        recom_form.escalation_recom_detail.render_kw = {"disabled": "disabled"}
        recom_form.action_detail.render_kw = {"disabled": "disabled"}
        recom_form.closure_detail.render_kw = {"disabled": "disabled"}
        recom_form.update_btn.render_kw = {"disabled": "disabled"}
        recom_form.close_btn.render_kw = {"disabled": "disabled"}
        
        if recom_info["Status"] == "Removed":
            recom_form.export_btn.render_kw = {"disabled": "disabled"}
            recom_form.remove_btn.render_kw = {"disabled": "disabled"}

    return render_template(
        'working_recom.html', recom_form=recom_form, case_id=case_id, recom_id=recom_id
        )
    

@recom_referral.route('/add_recom/<int:case_id>', methods=["GET", "POST"])
@login_required
def add_recom(case_id):
    add_recom_form = AddRecomendationForm()
    
    if add_recom_form.add_btn.data and add_recom_form.validate():
        sql_add = (
            "INSERT INTO Recommendation_local (Case_ID, Status, From_Section, Initiated_Date, Recomm_or_Escal, "
            "Followup_Date_1st, Followup_Date_2nd, Last_Followup_Date, "
            "Escalation_Date, Escalation_Type, Escalated_To, Recommendation_Closure_Details, "
            "Escal_Recomm_Details, Ack_Action_Date, Action_Details) VALUES "
            f"({case_id}, 'Open', ?, ?, ?, "
            "?, ?, ?, "
            "?, ?, ?, ?, "
            "?, ?, ?) ")
        insert_value_lst = (
            add_recom_form.from_section.data, datetime.now().date(), add_recom_form.recom_or_esc.data,
            add_recom_form.followup_1_date.data, add_recom_form.followup_2_date.data, add_recom_form.followup_last_date.data,
            add_recom_form.escalation_date.data, add_recom_form.escalation_type.data, add_recom_form.escalation_to.data, add_recom_form.closure_detail.data,
            add_recom_form.escalation_recom_detail.data, add_recom_form.ack_action_date.data, add_recom_form.action_detail.data)
        execute_sql(sql_add, insert_value_lst)
        
        sql_update = f"update CaseTracking_local SET Has_Recommendation = 'Yes', Escal_Recom_Status = 'Open' where Case_ID = {case_id}"
        execute_sql(sql_update)
        
        flash(f"Add a new Recommendation Record for Case {case_id}!", "success")
        return redirect(url_for('recom_referral.add_recom', case_id=case_id))

    return render_template(
        'add_recom.html', add_recom_form=add_recom_form, case_id=case_id
        )
     
   
#============================= Working on SAR Referral =====================================================  
@recom_referral.route('/sar/<int:case_id>', methods=["GET", "POST"])
@login_required
def preview_sar(case_id):
    filter_sar_form = FilterSAR()
    
    cursor, cnxn = connect_db()
    work_progress = pd.read_sql(f"select SAR_Referral_ID, Status from SARref_local where Case_ID = {case_id} and Status <> 'Removed' ", cnxn)
    total_tasks = work_progress["SAR_Referral_ID"].tolist()
    open_tasks = work_progress[work_progress["Status"]=="Open"]["SAR_Referral_ID"].tolist()
    cursor.close()
    cnxn.close()
    
    if filter_sar_form.add_sar.data and filter_sar_form.validate():
        return redirect(url_for("recom_referral.add_sar", case_id=case_id))
    
    if filter_sar_form.filter_sar.data and filter_sar_form.validate():
        cursor, cnxn = connect_db()
        sql_get_sars = (
            "select SAR_Referral_ID, Status, Subject_Name, Activity_Start_Date, Activity_End_Date, "
            "Amount, Subject_Account_NO, Referral_Reason, Reviewed_By_SAR, Referral_Necessary, SAR_Team_Comment, "
            "Referral_Warranted, EDD_Comment, Initiated_Date, CTRL_CMT, Date_Submitted, Date_Acknowledged "
            f"from SARref_local where Case_ID = {case_id} and Status <> 'Removed' ")
        if filter_sar_form.sar_status.data != "ALL":
             sql_get_sars += f" and Status = '{filter_sar_form.sar_status.data}' "
        sars = pd.read_sql(sql_get_sars, cnxn)
        cursor.close()
        cnxn.close()
        return render_template(
            'preview_case_sar.html', case_id=case_id, filter_sar_form=filter_sar_form, sars=sars,
            total_tasks=total_tasks, open_tasks=open_tasks
            )
    
    if filter_sar_form.export_sar.data and filter_sar_form.validate():        
        cursor, cnxn = connect_db()
        sql_get_sars = f'''select SAR_Referral_ID, Status, Subject_Name, Activity_Start_Date, Activity_End_Date, 
            Amount, Subject_Account_NO, Referral_Reason, Reviewed_By_SAR, Referral_Necessary, SAR_Team_Comment, 
            Referral_Warranted, EDD_Comment, Initiated_Date, CTRL_CMT, Date_Submitted, Date_Acknowledged 
            from SARref_local where Case_ID = {case_id} and Status <> 'Removed' '''
        if filter_sar_form.sar_status.data != "ALL":
             sql_get_sars += f" and Status = '{filter_sar_form.sar_status.data}' "        
        
        sar_df = pd.read_sql(sql_get_sars, cnxn)
        cursor.close()
        cnxn.close()
        
        # resp = make_response(sar_df.to_csv(index=False))
        # resp.headers["Content-Disposition"] = f"attachment; filename=SAR-CaseID_{case_id}-Status_{filter_sar_form.sar_status.data}.csv"
        # resp.headers["Content-Type"] = "text/csv"
        # return resp
    
        buffer = io.BytesIO()
        sar_df.to_excel(buffer, index=False)
        headers = {
            'Content-Disposition': f'attachment; filename=SAR-CaseID_{case_id}-Status_{filter_sar_form.sar_status.data}.xlsx',
            'Content-type': 'application/vnd.ms-excel'
        }
        return Response(buffer.getvalue(), mimetype='application/vnd.ms-excel', headers=headers)
    
    # cursor, cnxn = connect_db()
    # sql_get_sars = (
    #     "select SAR_Referral_ID, Status, Subject_Name, Activity_Start_Date, Activity_End_Date, "
    #     "Amount, Subject_Account_NO, Referral_Reason, Reviewed_By_SAR, Referral_Necessary, SAR_Team_Comment, "
    #     "Referral_Warranted, EDD_Comment, Initiated_Date, CTRL_CMT, Date_Submitted, Date_Acknowledged "
    #     f"from SARref_local where Case_ID = {case_id} and Status <> 'Removed' ")
    # sars = cursor.execute(sql_get_sars).fetchall()
    # cursor.close()
    # cnxn.close()
    
    cursor, cnxn = connect_db()
    sql_get_sars = f'''select SAR_Referral_ID, Status, Subject_Name, Activity_Start_Date, Activity_End_Date, 
        Amount, Subject_Account_NO, Referral_Reason, Reviewed_By_SAR, Referral_Necessary, SAR_Team_Comment, 
        Referral_Warranted, EDD_Comment, Initiated_Date, CTRL_CMT, Date_Submitted, Date_Acknowledged 
        from SARref_local where Case_ID = {case_id} and Status <> 'Removed' '''
    sars = pd.read_sql(sql_get_sars, cnxn)
    cursor.close()
    cnxn.close()
    
    return render_template(
        'preview_case_sar.html', case_id=case_id, filter_sar_form=filter_sar_form, sars=sars,
        total_tasks=total_tasks, open_tasks=open_tasks
        )


@recom_referral.route('/sar/<int:case_id>/<int:sar_id>', methods=["GET", "POST"])
@login_required
def each_sar(case_id, sar_id):
    sar_form = SARForm()
    
    if sar_form.remove_btn.data and sar_form.validate():
        sql_remove = f"update SARref_local set Status = 'Removed' where SAR_Referral_ID = {sar_id}"
        execute_sql(sql_remove)
        flash("Remove SAR Referral!", "warning")
        return redirect(url_for('recom_referral.each_sar', case_id=case_id, sar_id=sar_id))
    
    if sar_form.close_btn.data and sar_form.validate():
        sql_close = f"update SARref_local set Status = 'Closed', Date_Acknowledged = ? where SAR_Referral_ID = {sar_id}"
        execute_sql(sql_close, (datetime.now().date(),))
        flash("Close SAR Referral successfully!", "success")
        return redirect(url_for('recom_referral.each_sar', case_id=case_id, sar_id=sar_id))
    
    if sar_form.export_btn.data and sar_form.validate():
        # Export excel to user computer
        cursor, cnxn = connect_db()
        sql_get_sar = (
            "select Case_ID, SAR_Referral_ID, Status, Initiated_Date, Date_Submitted, Date_Acknowledged, Subject_Name, Amount, Subject_Account_NO, "
            "Referral_Reason, Reviewed_By_SAR, Referral_Necessary, SAR_Team_Comment, Referral_Warranted, EDD_Comment, CTRL_CMT, Status, "
            f"Currency, Activity_Start_Date, Activity_End_Date from SARref_local where SAR_Referral_ID = {sar_id}")
        sar_df = pd.read_sql(sql_get_sar, cnxn)
        cursor.close()
        cnxn.close()
        
        # resp = make_response(sar_df.to_csv(index=False))
        # resp.headers["Content-Disposition"] = f"attachment; filename=Case-{case_id}-SAR-{sar_id}.csv"
        # resp.headers["Content-Type"] = "text/csv"
        # return resp
    
        buffer = io.BytesIO()
        sar_df.to_excel(buffer, index=False)
        headers = {
            'Content-Disposition': f'attachment; filename=Case-{case_id}-SAR-{sar_id}.xlsx',
            'Content-type': 'application/vnd.ms-excel'
        }
        return Response(buffer.getvalue(), mimetype='application/vnd.ms-excel', headers=headers)
    
    if sar_form.update_btn.data and sar_form.validate():
        sql_update = f'''
            update SARref_local set Subject_Name = ?, Amount = ?, Subject_Account_NO = ?, Referral_Reason = ?, Reviewed_By_SAR = ?, 
            Referral_Necessary = ?, SAR_Team_Comment = ?, Referral_Warranted = ?, EDD_Comment = ?, CTRL_CMT = ?, Currency = ?, 
            Activity_Start_Date = ?, Activity_End_Date = ?, Date_Submitted = ? where SAR_Referral_ID = {sar_id} '''
        execute_sql(sql_update, (
            sar_form.subject_name.data, sar_form.amount.data, sar_form.subject_acc_no.data, sar_form.referral_reason.data, sar_form.reviewed_by_sar.data,
            sar_form.referral_necessary.data, sar_form.sar_team_comment.data, sar_form.referral_warranted.data, sar_form.edd_team_comment.data, sar_form.ctrl.data, 
            sar_form.currency.data, sar_form.activity_start_date.data, sar_form.activity_end_date.data, sar_form.date_submitted.data
        ))
        flash("Update SAR Referral successfully!", "success")
        return redirect(url_for('recom_referral.each_sar', case_id=case_id, sar_id=sar_id))
    
    sql_get_sar = f'''
        select Subject_Name, Amount, Subject_Account_NO, Referral_Reason, Reviewed_By_SAR, Referral_Necessary, SAR_Team_Comment, Referral_Warranted, 
        EDD_Comment, CTRL_CMT, Status, Currency, Activity_Start_Date, Activity_End_Date, Initiated_Date, Date_Submitted, Date_Acknowledged
        from SARref_local where SAR_Referral_ID = {sar_id}'''
    sar_info = select_sql(sql_get_sar)
    
    sar_form.sar_status.default = sar_info["Status"]
    sar_form.initiated_date.default = strptime_date(sar_info["Initiated_Date"])
    sar_form.date_acknowledged.default = strptime_date(sar_info["Date_Acknowledged"])
    sar_form.date_submitted.default = strptime_date(sar_info["Date_Submitted"])
    sar_form.subject_name.default = sar_info["Subject_Name"]
    sar_form.subject_acc_no.default = sar_info["Subject_Account_NO"]
    sar_form.amount.default = sar_info["Amount"]
    sar_form.currency.default = sar_info["Currency"]
    sar_form.activity_start_date.default = strptime_date(sar_info["Activity_Start_Date"])
    sar_form.activity_end_date.default = strptime_date(sar_info["Activity_End_Date"])
    sar_form.reviewed_by_sar.default = sar_info["Reviewed_By_SAR"]
    sar_form.referral_necessary.default = sar_info["Referral_Necessary"]
    sar_form.referral_warranted.default = sar_info["Referral_Warranted"]
    sar_form.ctrl.default = sar_info["CTRL_CMT"]
    sar_form.referral_reason.default = sar_info["Referral_Reason"]    
    sar_form.sar_team_comment.default = sar_info["SAR_Team_Comment"]    
    sar_form.edd_team_comment.default = sar_info["EDD_Comment"]    
    sar_form.process()
    
    # If validator is InputRequired(), then "disabled" will make form.validate() always False
    for myform_field in sar_form:
        validators = myform_field.validators
        for i, validator in enumerate(validators):
            if isinstance(validator, InputRequired):
                validators[i] = Optional()
                
    if sar_info["Status"] == "Removed" or sar_info["Status"] == "Closed":
        sar_form.date_submitted.render_kw = {"disabled": "disabled"}
        sar_form.subject_name.render_kw = {"disabled": "disabled"}
        sar_form.subject_acc_no.render_kw = {"disabled": "disabled"}
        sar_form.amount.render_kw = {"disabled": "disabled"}
        sar_form.currency.render_kw = {"disabled": "disabled"}
        sar_form.activity_start_date.render_kw = {"disabled": "disabled"}
        sar_form.activity_end_date.render_kw = {"disabled": "disabled"}
        sar_form.reviewed_by_sar.render_kw = {"disabled": "disabled"}
        sar_form.referral_necessary.render_kw = {"disabled": "disabled"}
        sar_form.referral_warranted.render_kw = {"disabled": "disabled"}
        sar_form.ctrl.render_kw = {"disabled": "disabled"}
        sar_form.referral_reason.render_kw = {"disabled": "disabled"}
        sar_form.sar_team_comment.render_kw = {"disabled": "disabled"}
        sar_form.edd_team_comment.render_kw = {"disabled": "disabled"}
        sar_form.update_btn.render_kw = {"disabled": "disabled"}
        sar_form.close_btn.render_kw = {"disabled": "disabled"}
        
        if sar_info["Status"] == "Removed":
            sar_form.export_btn.render_kw = {"disabled": "disabled"}
            sar_form.remove_btn.render_kw = {"disabled": "disabled"}

    return render_template(
        'working_sar.html', sar_form=sar_form, case_id=case_id, sar_id=sar_id
        )


@recom_referral.route('/add_sar/<int:case_id>', methods=["GET", "POST"])
@login_required
def add_sar(case_id):
    add_sar_form = AddSARForm()
    
    if add_sar_form.add_btn.data and add_sar_form.validate():
        sql_add = (
            f"INSERT INTO SARref_local (Case_ID, Status, Subject_Name, Subject_Account_NO, Amount, Activity_Start_Date, "
            "Activity_End_Date, Reviewed_By_SAR, Referral_Reason, SAR_Team_Comment, EDD_Comment) "
            f"VALUES ({case_id}, 'Open', ?, ?, ?, ?, ?, ?, ?, ?, ?) ")
        insert_value_lst = (
            add_sar_form.subject_name.data, add_sar_form.subject_acc_no.data, add_sar_form.amount.data,
            add_sar_form.activity_start_date.data, add_sar_form.activity_end_date.data, add_sar_form.reviewed_by_sar.data, add_sar_form.referral_reason.data, 
            add_sar_form.sar_team_comment.data, add_sar_form.edd_team_comment.data)
        execute_sql(sql_add, insert_value_lst)
        
        sql_update = f"update CaseTracking_local SET Has_SAR_Referral = 'Yes', SAR_Referral_Status = 'Open' where Case_ID = {case_id}"
        execute_sql(sql_update)
        
        flash(f"Add a new SAR Record for Case {case_id}!", "success")
        return redirect(url_for('recom_referral.add_sar', case_id=case_id))
        
    return render_template(
        'add_sar.html', add_sar_form=add_sar_form, case_id=case_id
        )
    
    
#============================= Working on Sanction Referral =====================================================  
@recom_referral.route('/sanction/<int:case_id>', methods=["GET", "POST"])
@login_required
def preview_sanc(case_id):
    filter_sanc_form = FilterSanc()
    
    cursor, cnxn = connect_db()
    work_progress = pd.read_sql(f"select Sanction_Referral_ID, Status from SanctionRef_local where Case_ID = {case_id} and Status <> 'Removed' ", cnxn)
    total_tasks = work_progress["Sanction_Referral_ID"].tolist()
    open_tasks = work_progress[work_progress["Status"]=="Open"]["Sanction_Referral_ID"].tolist()
    cursor.close()
    
    if filter_sanc_form.add_sanc.data and filter_sanc_form.validate():
        return redirect(url_for("recom_referral.add_sanc", case_id=case_id))
    
    if filter_sanc_form.filter_sanc.data and filter_sanc_form.validate():
        cursor, cnxn = connect_db()
        sql_get_sanctions = f'''
            select Case_ID, Sanction_Referral_ID, Status, Subject_Name, Amount, Currency, Referral_Reason, Additional_Comment, 
            Acknowledged_Date, Submit_Date from SanctionRef_local where Case_ID = {case_id} and Status <> 'Removed' '''
        if filter_sanc_form.sanc_status.data != "ALL":
             sql_get_sanctions += f" and Status = '{filter_sanc_form.sanc_status.data}' "
             
        sanctions = pd.read_sql(sql_get_sanctions, cnxn)
        cursor.close()
        cnxn.close()
        return render_template(
            'preview_case_sanc.html', case_id=case_id, filter_sanc_form=filter_sanc_form, sanctions=sanctions, 
            total_tasks=total_tasks, open_tasks=open_tasks
            )
    
    if filter_sanc_form.export_sanc.data and filter_sanc_form.validate():
        cursor, cnxn = connect_db()
        sql_get_sanctions = f'''
            select Case_ID, Sanction_Referral_ID, Status, Subject_Name, Amount, Currency, Referral_Reason, Additional_Comment, 
            Acknowledged_Date, Submit_Date from SanctionRef_local where Case_ID = {case_id} and Status <> 'Removed' '''
        if filter_sanc_form.sanc_status.data != "ALL":
             sql_get_sanctions += f" and Status = '{filter_sanc_form.sanc_status.data}' "        
        
        sanc_df = pd.read_sql(sql_get_sanctions, cnxn)
        cursor.close()
        cnxn.close()
        
        # resp = make_response(sanc_df.to_csv(index=False))
        # resp.headers["Content-Disposition"] = f"attachment; filename=Sanc-CaseID_{case_id}-Status_{filter_sanc_form.sanc_status.data}.csv"
        # resp.headers["Content-Type"] = "text/csv"
        # return resp
    
        buffer = io.BytesIO()
        sanc_df.to_excel(buffer, index=False)
        headers = {
            'Content-Disposition': f'attachment; filename=Sanc-CaseID_{case_id}-Status_{filter_sanc_form.sanc_status.data}.xlsx',
            'Content-type': 'application/vnd.ms-excel'
        }
        return Response(buffer.getvalue(), mimetype='application/vnd.ms-excel', headers=headers)
    
    cursor, cnxn = connect_db()
    sql_get_sanctions = f'''
        select Case_ID, Sanction_Referral_ID, Status, Subject_Name, Amount, Currency, Referral_Reason, Additional_Comment, 
        Acknowledged_Date, Submit_Date from SanctionRef_local where Case_ID = {case_id} and Status <> 'Removed' '''
    sanctions = pd.read_sql(sql_get_sanctions, cnxn)
    cursor.close()
    cnxn.close()
    
    return render_template(
        'preview_case_sanc.html', case_id=case_id, filter_sanc_form=filter_sanc_form, sanctions=sanctions, 
        total_tasks=total_tasks, open_tasks=open_tasks
        )


@recom_referral.route('/sanction/<int:case_id>/<int:sanc_id>', methods=["GET", "POST"])
@login_required
def each_sanc(case_id, sanc_id):
    sanc_form = SanctionForm()
    sanc_form.date_acknowledged.render_kw = {"disabled": "disabled"}
    
    if sanc_form.remove_btn.data and sanc_form.validate():
        sql_remove = f"update SanctionRef_local set Status = 'Removed' where Sanction_Referral_ID = {sanc_id}"
        execute_sql(sql_remove)
        flash("Remove Sanction Referral!", "warning")
        return redirect(url_for('recom_referral.each_sanc', case_id=case_id, sanc_id=sanc_id))

    if sanc_form.close_btn.data and sanc_form.validate():
        sql_close = f"update SanctionRef_local set Status = 'Closed', Acknowledged_Date = ? where Sanction_Referral_ID = {sanc_id}"
        execute_sql(sql_close, (datetime.now().date(),))
        flash("Close Sanction Referral successfully!", "success")
        return redirect(url_for('recom_referral.each_sanc', case_id=case_id, sanc_id=sanc_id))
    
    if sanc_form.export_btn.data and sanc_form.validate():
        # Export excel to user computer
        cursor, cnxn = connect_db()
        sql_get_sanc = (
            "select Case_ID, Sanction_Referral_ID, Status, Acknowledged_Date, Subject_Name, Amount, Referral_Reason, Additional_Comment, Submit_Date "
            f"from SanctionRef_local where Sanction_Referral_ID = {sanc_id}")
        sanc_df = pd.read_sql(sql_get_sanc, cnxn)
        cursor.close()
        cnxn.close()
        
        # resp = make_response(sanc_df.to_csv(index=False))
        # resp.headers["Content-Disposition"] = f"attachment; filename=Case-{case_id}-Sanc-{sanc_id}.csv"
        # resp.headers["Content-Type"] = "text/csv"
        # return resp
    
        buffer = io.BytesIO()
        sanc_df.to_excel(buffer, index=False)
        headers = {
            'Content-Disposition': f'attachment; filename=Case-{case_id}-Sanc-{sanc_id}.xlsx',
            'Content-type': 'application/vnd.ms-excel'
        }
        return Response(buffer.getvalue(), mimetype='application/vnd.ms-excel', headers=headers)
    
    if sanc_form.update_btn.data and sanc_form.validate():
        sql_update = (
            "update SanctionRef_local set Subject_Name = ?, Amount = ?, Referral_Reason = ?, Additional_Comment = ?, Submit_Date = ? "
            f"where Sanction_Referral_ID = {sanc_id}")
        execute_sql(sql_update, (
            sanc_form.subject_name.data, sanc_form.amount.data, sanc_form.referral_reason.data, sanc_form.additional_comment.data, sanc_form.date_submitted.data
        ))
        flash("Update Sanction Referral successfully!", "success")
        return redirect(url_for('recom_referral.each_sanc', case_id=case_id, sanc_id=sanc_id))
    
    sql_get_sanc = f'''
        select Status, Subject_Name, Amount, Currency, Referral_Reason, Additional_Comment, Submit_Date, Acknowledged_Date 
        from SanctionRef_local where Sanction_Referral_ID = {sanc_id}'''    
    sanc_info = select_sql(sql_get_sanc)
    
    sanc_form.date_acknowledged.default = strptime_date(sanc_info["Acknowledged_Date"])
    sanc_form.date_submitted.default = strptime_date(sanc_info["Submit_Date"])
    sanc_form.subject_name.default = sanc_info["Subject_Name"]
    sanc_form.amount.default = sanc_info["Amount"]
    sanc_form.currency.default = sanc_info["Currency"]
    sanc_form.date_acknowledged.default = strptime_date(sanc_info["Acknowledged_Date"])
    sanc_form.referral_reason.default = sanc_info["Referral_Reason"]
    sanc_form.additional_comment.default = sanc_info["Additional_Comment"]
    sanc_form.process()
    
    # If validator is InputRequired(), then "disabled" will make form.validate() always False
    for myform_field in sanc_form:
        validators = myform_field.validators
        for i, validator in enumerate(validators):
            if isinstance(validator, InputRequired):
                validators[i] = Optional()
    
    if sanc_info["Status"] == "Removed" or sanc_info["Status"] == "Closed":
        sanc_form.date_submitted.render_kw = {"disabled": "disabled"}
        sanc_form.subject_name.render_kw = {"disabled": "disabled"}
        sanc_form.amount.render_kw = {"disabled": "disabled"}
        sanc_form.currency.render_kw = {"disabled": "disabled"}
        sanc_form.date_acknowledged.render_kw = {"disabled": "disabled"}
        sanc_form.referral_reason.render_kw = {"disabled": "disabled"}
        sanc_form.additional_comment.render_kw = {"disabled": "disabled"}
        sanc_form.update_btn.render_kw = {"disabled": "disabled"}
        sanc_form.close_btn.render_kw = {"disabled": "disabled"}
        
        if sanc_info["Status"] == "Removed":
            sanc_form.export_btn.render_kw = {"disabled": "disabled"}
            sanc_form.remove_btn.render_kw = {"disabled": "disabled"}

    return render_template(
        'working_sanc.html', sanc_form=sanc_form, case_id=case_id, sanc_id=sanc_id
        )
    
    
@recom_referral.route('/add_sanction/<int:case_id>', methods=["GET", "POST"])
@login_required
def add_sanc(case_id):
    add_sanc_form = AddSanctionForm()
    
    if add_sanc_form.add_btn.data and add_sanc_form.validate():
        sql_add = (
            f"INSERT INTO SanctionRef_local (Case_ID, Status, Subject_Name, Amount, Referral_Reason, Additional_Comment, Submit_Date) "
            f"VALUES ({case_id}, 'Open', ?, ?, ?, ?, ?) ")
        insert_value_lst = (
            add_sanc_form.subject_name.data, add_sanc_form.amount.data, add_sanc_form.referral_reason.data,
            add_sanc_form.additional_comment.data, add_sanc_form.date_submitted.data)
        execute_sql(sql_add, insert_value_lst)
        
        sql_update = f"update CaseTracking_local SET Has_Sanction_Referral = 'Yes', Sanction_Referral_Status = 'Open' where Case_ID = {case_id}"
        execute_sql(sql_update)
        
        flash(f"Add a new Sanction Record for Case {case_id}!", "success")
        return redirect(url_for('recom_referral.add_sanc', case_id=case_id))
        
    return render_template(
        'add_sanc.html', add_sanc_form=add_sanc_form, case_id=case_id
        )
